#!/usr/bin/env python3
"""Generate a Word user manual for Andes Padel admin panel (Spanish).

Based on actual codebase:
- apps/adminpanel/views.py (all views and POST actions)
- apps/adminpanel/templates/ (all templates)
- apps/users/models.py (User model: roles, statuses)
- apps/bookings/models.py (Booking states)
- apps/payments/models.py (Payment states, methods)
- apps/scheduling/models.py (TimeSlot states)
- apps/reports/services.py (ReportService)
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

IMAGES = Path(__file__).parent
OUTPUT = IMAGES.parent / "Manual_Usuario_AndesPadel.docx"

doc = Document()

# ── Styles ──
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x00, 0x2F, 0x48)  # Midnight Blue


# ── Helper functions ──
def add_screenshot(name, caption, width=5.5):
    path = IMAGES / f"{name}.png"
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.style = doc.styles["Normal"]
        run2 = cap.runs[0] if cap.runs else cap.add_run()
        run2.font.size = Pt(9)
        run2.font.italic = True
        run2.font.color.rgb = RGBColor(0x5B, 0x64, 0x72)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f" {text}")
    else:
        p.add_run(text)


# ════════════════════════════════════════════════════════════════════════
#  PORTADA
# ════════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Manual de Usuario")
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x2F, 0x48)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Andes Pádel — Sistema de Gestión de Reservas")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("Versión 1.0 — Agosto 2026\n")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run = meta.add_run("Plataforma: Panel de Administración Web\n")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run = meta.add_run("Documento confidencial — Solo para uso interno")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  ÍNDICE
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("Índice", level=1)

toc_items = [
    "1. Introducción",
    "2. Acceso al Sistema y Roles",
    "3. Panel de Administración Django (/admin/)",
    "   3.1. Panel Principal (Dashboard)",
    "   3.2. Gestión de Usuarios",
    "   3.3. Gestión de Canchas y Sedes",
    "   3.4. Gestión de Reservas",
    "   3.5. Gestión de Eventos, Torneos y Noticias",
    "   3.6. Gestión de Horarios (TimeSlots)",
    "   3.7. Gestión de Reglas de Precios",
    "   3.8. Gestión de Pagos",
    "   3.9. Gestión de Notificaciones",
    "   3.10. Políticas de Cancelación",
    "   3.11. Registro de Auditoría",
    "4. Panel de Administración Personalizado (/adminpanel/)",
    "   4.1. Panel de Control (Dashboard)",
    "   4.2. Calendario de Reservas",
    "   4.3. Gestión de Canchas y Mantenimiento",
    "   4.4. Gestión de Usuarios y Roles (RBAC)",
    "   4.5. Pagos y Verificación de Transferencias",
    "   4.6. Torneos, Eventos y Noticias",
    "   4.7. Analítica de Negocio y Reportes",
    "   4.8. Configuración (Políticas y Tarifas)",
    "   4.9. Registro de Auditoría y Seguridad",
    "5. Documentación de la API (Swagger / ReDoc)",
    "   5.1. Endpoints Disponibles",
    "   5.2. Autenticación JWT",
    "6. Cuentas de Prueba",
    "7. Solución de Problemas",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    if not item.startswith("   "):
        for r in p.runs:
            r.font.bold = True

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  1. INTRODUCCIÓN
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("1. Introducción", level=1)

doc.add_paragraph(
    "Andes Pádel es un sistema completo de gestión de reservas de canchas de pádel "
    "diseñado para el complejo deportivo Andes Pádel en Quito, Ecuador. El sistema "
    "comprende una API RESTful (Django 5.2 + DRF), una aplicación móvil Flutter y "
    "dos paneles de administración web."
)
doc.add_paragraph(
    "Este manual describe paso a paso cómo utilizar los paneles de administración "
    "web para gestionar usuarios, canchas, reservas, pagos, eventos, torneos, "
    "reportes y configuración del complejo deportivo."
)

doc.add_heading("Arquitectura del Sistema", level=2)
add_table(
    ["Componente", "Tecnología", "URL / Acceso"],
    [
        ["API REST", "Django 5.2 + DRF", "http://[IP]:8000/api/"],
        ["Documentación API", "drf-spectacular", "http://[IP]:8000/api/docs/"],
        ["Admin Django", "Django Admin + RBAC", "http://[IP]:8000/admin/"],
        ["Admin Personalizado", "Django Templates", "http://[IP]:8000/adminpanel/"],
        ["App Móvil", "Flutter 3.27", "APK Android/iOS"],
        ["Tareas Asíncronas", "Celery + Redis", "Worker en background"],
    ],
)

doc.add_heading("Paleta Corporativa", level=2)
doc.add_paragraph(
    "El sistema utiliza la identidad visual corporativa de Andes Pádel:"
)
add_table(
    ["Color", "Nombre", "Hex", "Uso"],
    [
        ["Primario", "Midnight Blue", "#002F48", "Botones principales, sidebar, encabezados"],
        ["Secundario", "Celeste", "#3571B8", "Hover states, enlaces"],
        ["Acento", "Verde Limón", "#CEDC29", "Badges, highlights, indicadores"],
        ["Acento Suave", "Light Verde Limón", "#E8EFB0", "Fondos suaves, chips"],
        ["Éxito", "Green", "#2E7D32", "Estados confirmados"],
        ["Advertencia", "Amber", "#FFC107", "Estados pendientes"],
        ["Peligro", "Red", "#D32F2F", "Estados cancelados, errores"],
    ],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  2. ACCESO AL SISTEMA Y ROLES
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Acceso al Sistema y Roles", level=1)

doc.add_paragraph(
    "El sistema cuenta con dos paneles de administración web, ambos accesibles "
    "desde el navegador. Utilizan las mismas credenciales de usuario."
)

doc.add_heading("Panel Personalizado (Recomendado para operación diaria)", level=2)
doc.add_paragraph(
    "1. Navegue a http://[IP_DEL_SERVIDOR]:8000/adminpanel/\n"
    "2. Será redirigido a la pantalla de inicio de sesión\n"
    "3. Ingrese su correo electrónico y contraseña\n"
    "4. Haga clic en \"Iniciar sesión\""
)
doc.add_paragraph(
    "Este panel requiere un rol de personal (recepcionista, gerente, dueño o superadmin). "
    "Los usuarios con rol de cliente no pueden acceder."
)

doc.add_heading("Panel Django Admin (Gestión técnica avanzada)", level=2)
doc.add_paragraph(
    "1. Navegue a http://[IP_DEL_SERVIDOR]:8000/admin/\n"
    "2. Ingrese su correo electrónico en el campo \"Email\"\n"
    "3. Ingrese su contraseña en el campo \"Contraseña\"\n"
    "4. Haga clic en \"Iniciar sesión\""
)

doc.add_heading("Sistema de Roles (RBAC)", level=2)
doc.add_paragraph(
    "El acceso está controlado por un sistema de roles jerárquicos. "
    "Cada rol tiene permisos específicos en cada sección del sistema."
)

add_table(
    ["Rol", "Acceso Panel", "Acceso Django Admin", "Descripción"],
    [
        ["superadmin", "Completo", "Completo", "Acceso total a todos los módulos y configuración"],
        ["dueño", "Completo", "Completo (excepto finanzas edit)", "Propietario, acceso amplio"],
        ["gerente", "Completo", "CRUD no-financiero, vista financiero", "Gestión operativa diaria"],
        ["recepcionista", "Completo", "Solo lectura/edición no-financiero", "Atención al cliente, reservas"],
        ["cliente", "Sin acceso", "Sin acceso", "Solo app móvil"],
    ],
)

doc.add_paragraph()
doc.add_paragraph(
    "Nota: Los permisos financieros (Payment, CancellationPolicy) están restringidos "
    "a superadmin y dueño. El gerente puede ver pero no editar estos módulos."
)

add_screenshot("03_admin_login", "Figura 1: Pantalla de inicio de sesión del panel")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  3. PANEL DE ADMINISTRACIÓN DJANGO
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Panel de Administración Django (/admin/)", level=1)

doc.add_paragraph(
    "El panel de administración Django proporciona acceso completo a todos los "
    "modelos del sistema con control de acceso basado en roles (RBAC). "
    "Es la herramienta principal para la gestión técnica de los datos."
)

# 3.1 Dashboard
doc.add_heading("3.1. Panel Principal (Dashboard)", level=2)
doc.add_paragraph(
    "Al iniciar sesión, se muestra el panel principal con un resumen de todas "
    "las secciones disponibles agrupadas por categoría:"
)

doc.add_heading("Secciones disponibles:", level=3)
add_bullet("Autenticación y Usuarios — Gestión de cuentas, roles y permisos")
add_bullet("Canchas y Sedes — Administración de canchas, horarios y venues")
add_bullet("Reservas — Lista completa de reservas con estados y acciones")
add_bullet("Pagos — Registro y seguimiento de transacciones (restringido)")
add_bullet("Eventos — Torneos, eventos especiales y noticias")
add_bullet("Programación — TimeSlots, BookingHolds, mantenimientos")
add_bullet("Notificaciones — Sistema de notificaciones in-app y preferencias")
add_bullet("Seguridad — Registro de auditoría completo (solo lectura)")
add_bullet("Precios — Reglas de pricing y feriados")
add_bullet("Políticas — Reglas de cancelación y penalidades (restringido)")

add_screenshot("04_admin_dashboard", "Figura 2: Panel principal de administración Django")

# 3.2 Usuarios
doc.add_heading("3.2. Gestión de Usuarios", level=2)
doc.add_paragraph(
    "La sección de Usuarios permite administrar todas las cuentas del sistema. "
    "Cada usuario incluye perfil completo con roles y control de acceso."
)

doc.add_heading("Campos del usuario:", level=3)
add_bullet("Correo electrónico (identificador único)")
add_bullet("Nombre completo")
add_bullet("Teléfono")
add_bullet("Rol: superadmin, dueño, gerente, recepcionista, cliente")
add_bullet("Estado: active, suspended, blocked")
add_bullet("Verificación de email (email_verified)")
add_bullet("Idioma preferido (es/en/pt/ca)")
add_bullet("Consentimiento GDPR (consent_version, consent_granted)")
add_bullet("Fecha de registro (date_joined)")

add_screenshot("05_admin_users", "Figura 3: Gestión de usuarios")

# 3.3 Canchas
doc.add_heading("3.3. Gestión de Canchas y Sedes", level=2)
doc.add_paragraph(
    "Permite administrar las sedes (venues) y canchas de pádel del complejo. "
    "Cada cancha pertenece a una sede y tiene propiedades específicas."
)

doc.add_heading("Sedes (Venues):", level=3)
add_bullet("Nombre del complejo (ej: Andes Pádel Club)")
add_bullet("Dirección física")
add_bullet("Zona horaria y moneda")
add_bullet("Estado: activo/inactivo")

doc.add_heading("Canchas:", level=3)
add_bullet("Nombre (ej: C1, C2)")
add_bullet("Tipo de cancha: techada o abierta")
add_bullet("Iluminación LED disponible (sí/no)")
add_bullet("Precio base por hora")
add_bullet("Estado: activa o inactiva")
add_bullet("Horarios de la cancha (CourtSchedule) — apertura/cierre por día")

add_screenshot("06_admin_courts", "Figura 4: Gestión de canchas")

# 3.4 Reservas
doc.add_heading("3.4. Gestión de Reservas", level=2)
doc.add_paragraph(
    "Muestra todas las reservas realizadas en el sistema. El admin permite "
    "acciones masivas como confirmar o cancelar múltiples reservas."
)

doc.add_heading("Estados de una reserva:", level=3)
add_table(
    ["Estado", "Descripción"],
    [
        ["pending_payment", "Reserva creada, esperando pago"],
        ["confirmed", "Reserva confirmada y pagada"],
        ["in_progress", "Reserva en curso (fecha/hora actual)"],
        ["completed", "Reserva finalizada"],
        ["cancelled", "Reserva cancelada"],
        ["no_show", "El cliente no se presentó"],
    ],
)
doc.add_paragraph()

doc.add_heading("Información visible por cada reserva:", level=3)
add_bullet("Usuario que realizó la reserva (email)")
add_bullet("Cancha asignada")
add_bullet("Fecha y hora de inicio/fin")
add_bullet("Duración en minutos")
add_bullet("Número de jugadores")
add_bullet("Precio total")
add_bullet("Estado actual (badge de color)")
add_bullet("Slots horarios reservados (BookingSlot inline)")
add_bullet("Historial de cambios de estado (BookingEvent inline)")

doc.add_heading("Acciones disponibles:", level=3)
add_bullet("Confirmar reservas — Cambia estado a 'confirmed'")
add_bullet("Cancelar reservas — Cambia estado a 'cancelled'")

add_screenshot("07_admin_bookings", "Figura 5: Gestión de reservas")

# 3.5 Eventos
doc.add_heading("3.5. Gestión de Eventos, Torneos y Noticias", level=2)
doc.add_paragraph(
    "Permite crear y administrar torneos, eventos especiales y noticias "
    "para los usuarios del complejo."
)

doc.add_heading("Torneos:", level=3)
add_bullet("Nombre y categoría (Open, Intermedio, etc.)")
add_bullet("Máximo de parejas/equipos")
add_bullet("Cuota de inscripción ($)")
add_bullet("Fecha de inicio y fin")
add_bullet("Fecha límite de inscripción")
add_bullet("Estado: open, in_progress, completed, cancelled")

doc.add_heading("Eventos:", level=3)
add_bullet("Título y descripción")
add_bullet("Fecha/hora de inicio y fin")
add_bullet("Ubicación")
add_bullet("Estado: draft, published, cancelled")

doc.add_heading("Noticias:", level=3)
add_bullet("Título y contenido")
add_bullet("Fecha de publicación")
add_bullet("Estado: draft, published")

add_screenshot("08_admin_events", "Figura 6: Gestión de eventos y torneos")

# 3.6 Horarios
doc.add_heading("3.6. Gestión de Horarios (TimeSlots)", level=2)
doc.add_paragraph(
    "Las franjas horarias (TimeSlots) representan los bloques de 30 minutos "
    "disponibles para reserva en cada cancha. Este sistema permite un control "
    "granular de la disponibilidad."
)

doc.add_heading("Estados de un horario:", level=3)
add_table(
    ["Estado", "Descripción"],
    [
        ["available", "Disponible para reserva"],
        ["held", "Mantenido temporalmente (carrito de compra)"],
        ["booked", "Reservado por un cliente"],
        ["blocked", "Bloqueado por administración"],
    ],
)
doc.add_paragraph()

doc.add_paragraph(
    "Nota: Los TimeSlots no se crean manualmente. Se generan automáticamente "
    "mediante el comando 'generate_timeslots' basado en los CourtSchedule de cada cancha."
)

add_screenshot("09_admin_timeslots", "Figura 7: Gestión de horarios")

# 3.7 Precios
doc.add_heading("3.7. Gestión de Reglas de Precios", level=2)
doc.add_paragraph(
    "Las reglas de precios permiten definir tarifas dinámicas según la cancha, "
    "el día de la semana, la hora y otros factores."
)

doc.add_heading("Parámetros de cada regla:", level=3)
add_bullet("Nombre descriptivo de la regla")
add_bullet("Sede aplicable")
add_bullet("Zona horaria (si aplica)")
add_bullet("Tipo de cancha (techada/abierta)")
add_bullet("Día de la semana (0=Lunes - 6=Domingo)")
add_bullet("Hora de inicio y fin")
add_bullet("Multiplicador de precio (ej: 1.5 = 50% más caro)")
add_bullet("Prioridad de la regla (para resolución de conflictos)")

doc.add_heading("Feriados:", level=3)
add_bullet("Fecha del feriado")
add_bullet("Nombre del feriado")
add_bullet("Sede aplicable")

add_screenshot("10_admin_pricing", "Figura 8: Gestión de reglas de precios")

# 3.8 Pagos
doc.add_heading("3.8. Gestión de Pagos", level=2)
doc.add_paragraph(
    "Registra y administra todos los pagos asociados a reservas. "
    "Este módulo está restringido a superadmin y dueño (FINANCIAL=True)."
)

doc.add_heading("Estados de pago:", level=3)
add_table(
    ["Estado", "Descripción"],
    [
        ["pending", "Pago pendiente de procesamiento"],
        ["pending_transfer", "Esperando confirmación de transferencia"],
        ["captured", "Pago capturado exitosamente (Stripe)"],
        ["confirmed", "Pago confirmado manualmente"],
        ["failed", "Pago fallido o rechazado"],
        ["refunded", "Reembolso procesado"],
    ],
)
doc.add_paragraph()

doc.add_heading("Métodos de pago soportados:", level=3)
add_bullet("stripe — Tarjeta de crédito/débito vía Stripe")
add_bullet("transfer — Transferencia bancaria (requiere confirmación manual)")
add_bullet("cash — Pago en efectivo")

doc.add_paragraph(
    "Restricciones por rol:\n"
    "- superadmin / dueño: acceso completo\n"
    "- gerente: solo lectura (no puede crear/editar pagos)\n"
    "- recepcionista: bloqueado completamente"
)

add_screenshot("11_admin_payments", "Figura 9: Gestión de pagos")

# 3.9 Notificaciones
doc.add_heading("3.9. Gestión de Notificaciones", level=2)
doc.add_paragraph(
    "El sistema de notificaciones permite enviar avisos a los usuarios "
    "a través de múltiples canales: in-app, email y push (FCM)."
)

doc.add_heading("Tipos de notificación:", level=3)
add_bullet("booking_confirmed — Confirmación de reserva")
add_bullet("booking_cancelled — Cancelación de reserva")
add_bullet("booking_reminder — Recordatorio de reserva")
add_bullet("payment_received — Pago recibido")
add_bullet("tournament_registration — Inscripción a torneo")
add_bullet("news_published — Nueva noticia publicada")

doc.add_heading("Preferencias de notificación:", level=3)
add_bullet("Cada usuario puede configurar qué tipos de notificación recibir")
add_bullet("Canales: in_app, email, push (FCM)")
add_bullet("Las preferencias se gestionan desde NotificationPreference")

doc.add_paragraph(
    "Nota: Las notificaciones in-app y por email funcionan sin configuración adicional. "
    "Las notificaciones push requieren configuración de Firebase Cloud Messaging."
)

add_screenshot("12_admin_notifications", "Figura 10: Gestión de notificaciones")

# 3.10 Políticas
doc.add_heading("3.10. Políticas de Cancelación", level=2)
doc.add_paragraph(
    "Las políticas de cancelación definen las reglas y penalidades cuando "
    "un usuario cancela una reserva. Este módulo está restringido (FINANCIAL=True)."
)

doc.add_heading("Parámetros configurables:", level=3)
add_bullet("Nombre de la política")
add_bullet("Horas de antelación mínima para cancelar sin penalidad (free_cancellation_hours)")
add_bullet("Porcentaje de penalidad por cancelación tardía (late_cancellation_fee_percent)")
add_bullet("Porcentaje de penalidad por no-show (no_show_fee_percent)")
add_bullet("Tiempo de hold en minutos (hold_duration_minutes)")
add_bullet("Descripción de la política")

add_screenshot("14_admin_policies", "Figura 11: Políticas de cancelación")

# 3.11 Auditoría
doc.add_heading("3.11. Registro de Auditoría", level=2)
doc.add_paragraph(
    "El registro de auditoría captura automáticamente todas las acciones "
    "realizadas en el sistema. Es completamente de solo lectura — no se pueden "
    "crear, editar ni eliminar registros."
)

doc.add_heading("Información registrada:", level=3)
add_bullet("Usuario que realizó la acción (user email)")
add_bullet("Tipo de acción (action): login, booking_create, court_toggle, etc.")
add_bullet("Entidad afectada (entity): user, booking, court, payment, etc.")
add_bullet("ID de la entidad (entity_id)")
add_bullet("Fecha y hora exacta (created_at)")
add_bullet("Dirección IP (ip_address)")

add_screenshot("15_admin_audit_logs", "Figura 12: Registro de auditoría")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  4. PANEL PERSONALIZADO
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Panel de Administración Personalizado (/adminpanel/)", level=1)

doc.add_paragraph(
    "El panel personalizado ofrece una interfaz optimizada para las "
    "operaciones diarias del complejo deportivo. A diferencia del admin Django, "
    "está diseñado para tareas frecuentes como crear reservas, bloquear horarios "
    "y verificar pagos con una experiencia visual moderna."
)

doc.add_paragraph(
    "Requisito: Solo usuarios con rol de personal (recepcionista, gerente, "
    "dueño o superadmin) pueden acceder. El control de acceso se verifica "
    "automáticamente en cada vista."
)

# 4.1 Dashboard
doc.add_heading("4.1. Panel de Control (Dashboard)", level=2)
doc.add_paragraph(
    "El dashboard principal muestra un resumen ejecutivo del estado actual "
    "del complejo deportivo con métricas en tiempo real."
)

doc.add_heading("KPIs en tarjetas:", level=3)
add_bullet("Ingresos de Hoy — Total de pagos recibidos el día actual ($)")
add_bullet("Reservas Activas — Cantidad de reservas con estado confirmed/in_progress")
add_bullet("Ocupación de Canchas — Porcentaje de slots reservados vs disponibles")
add_bullet("Total Usuarios — Número total de usuarios registrados")

doc.add_heading("Próximas Reservas:", level=3)
add_bullet("Tabla con las 10 reservas del día actual")
add_bullet("Columnas: Cliente, Cancha, Hora Inicio, Estado (badge)")

doc.add_heading("Alertas del Sistema:", level=3)
add_bullet("Transferencias Pendientes — Pagos esperando confirmación")
add_bullet("Reservas Sin Pagar — Reservas con estado pending_payment")
add_bullet("Mantenimientos Hoy — Canchas con mantenimiento programado")

doc.add_heading("Accesos Rápidos:", level=3)
add_bullet("Ver Calendario Full → Navega al calendario de reservas")
add_bullet("Revisar Transferencias → Filtra pagos pendientes")

add_screenshot("16_custom_admin_dashboard", "Figura 13: Panel de control principal")

# 4.2 Calendario
doc.add_heading("4.2. Calendario de Reservas", level=2)
doc.add_paragraph(
    "El calendario es la herramienta principal para la gestión diaria. Muestra "
    "una cuadrícula visual de todos los horarios por cancha, permitiendo "
    "bloquear, desbloquear y crear reservas directamente."
)

doc.add_heading("Navegación:", level=3)
add_bullet("Botones ← → para navegar entre días")
add_bullet("Fecha actual visible en el encabezado")
add_bullet("Parámetro URL: ?date=YYYY-MM-DD")

doc.add_heading("Vista de cuadrícula:", level=3)
add_bullet("Filas = Franjas horarias de 30 minutos (06:00 a 24:00)")
add_bullet("Columnas = Canchas activas")
add_bullet("Cada celda muestra: estado + usuario (si está reservado)")

doc.add_heading("Acciones por slot:", level=3)
add_table(
    ["Estado Actual", "Acción", "Efecto"],
    [
        ["available", "Bloquear", "Cambia a blocked (motivo obligatorio)"],
        ["blocked", "Liberar", "Cambia a available"],
        ["booked", "—", "Muestra email del cliente"],
        ["held", "—", "Muestra temporalidad de hold"],
    ],
)
doc.add_paragraph()

doc.add_heading("Crear Reserva Manual:", level=3)
doc.add_paragraph(
    "1. Seleccione la fecha en el calendario\n"
    "2. En la sección \"Crear reserva manual\", seleccione la cancha\n"
    "3. Seleccione el cliente (desplegable con todos los usuarios activos)\n"
    "4. Ingrese la hora de inicio (formato HH:MM)\n"
    "5. Seleccione la duración (60, 90 o 120 minutos)\n"
    "6. Haga clic en \"Crear reserva\"\n"
    "7. Se creará la reserva con estado \"confirmada\" y precio fijo $30.00\n"
    "8. Se enviará una notificación al cliente vía Celery"
)

doc.add_paragraph(
    "Nota: La reserva manual se crea directamente con estado 'confirmed' "
    "y un precio fijo de $30.00. Las notificaciones se envían "
    "asíncronamente mediante Celery."
)

add_screenshot("17_custom_admin_calendar", "Figura 14: Calendario de reservas")

# 4.3 Canchas Panel
doc.add_heading("4.3. Gestión de Canchas y Mantenimiento", level=2)
doc.add_paragraph(
    "Vista simplificada para la gestión rápida de canchas con acciones frecuentes."
)

doc.add_heading("Tabla de Canchas:", level=3)
add_bullet("Columnas: Nombre, Tipo (badge), Iluminación (✓ LED / No), Estado, Acción")
add_bullet("Botón de toggle: Activar/Desactivar cancha con un clic")

doc.add_heading("Acciones disponibles:", level=3)
add_bullet("Activar/Desactivar cancha — Cambia el estado active/inactive")
add_bullet("Crear nueva cancha — Nombre, tipo (techada/abierta), iluminación LED")
add_bullet("Programar mantenimiento — Cancha, motivo, fecha/hora inicio-fin")

doc.add_heading("Tabla de Mantenimientos:", level=3)
add_bullet("Muestra los últimos 20 mantenimientos programados")
add_bullet("Columnas: Cancha, Motivo, Inicio, Fin")

add_screenshot("18_custom_admin_courts", "Figura 15: Gestión de canchas y mantenimiento")

# 4.4 Usuarios Panel
doc.add_heading("4.4. Gestión de Usuarios y Roles (RBAC)", level=2)
doc.add_paragraph(
    "Lista paginada de usuarios con opciones de búsqueda, filtrado y "
    "cambio rápido de roles y estados."
)

doc.add_heading("Filtros disponibles:", level=3)
add_bullet("Búsqueda — Por nombre o correo electrónico (icontains)")
add_bullet("Filtrar por rol — recepcionista, gerente, dueño, superadmin")
add_bullet("Filtrar por estado — active, suspended, blocked")

doc.add_heading("Columnas de la tabla:", level=3)
add_bullet("Usuario (email)")
add_bullet("Nombre completo")
add_bullet("Rol Actual (badge de color)")
add_bullet("Estado (badge: active=verde, suspended=amarillo, blocked=rojo)")
add_bullet("Fecha de registro")
add_bullet("Cambiar Rol (dropdown + botón Guardar)")
add_bullet("Cambiar Estado (dropdown + botón Guardar)")

doc.add_paragraph(
    "Acciones:\n"
    "- change_role: Actualiza el rol del usuario (audit: admin.user_role_change)\n"
    "- change_status: Actualiza el estado del usuario (audit: admin.user_status_change)"
)

doc.add_paragraph("Paginación: 30 usuarios por página.")

add_screenshot("19_custom_admin_users", "Figura 16: Gestión de usuarios con RBAC")

# 4.5 Pagos Panel
doc.add_heading("4.5. Pagos y Verificación de Transferencias", level=2)
doc.add_paragraph(
    "Interfaz para verificar y gestionar pagos, especialmente útil para "
    "confirmar transferencias bancarias."
)

doc.add_heading("Filtros disponibles:", level=3)
add_bullet("Por estado del pago (pending_transfer, confirmed, captured, refunded, failed)")
add_bullet("Por método de pago (transfer, stripe, cash)")

doc.add_heading("Columnas de la tabla:", level=3)
add_bullet("ID Pago (UUID truncado)")
add_bullet("Cliente (email)")
add_bullet("Monto ($ en verde)")
add_bullet("Método (badge)")
add_bullet("Estado (badge)")
add_bullet("Comprobante / Referencia")
add_bullet("Fecha")
add_bullet("Acciones (botones según estado)")

doc.add_heading("Acciones por estado:", level=3)
add_table(
    ["Estado", "Acciones Disponibles"],
    [
        ["pending_transfer", "Aprobar (confirm_transfer), Ver comprobante, Rechazar (con motivo)"],
        ["confirmed / captured", "Reembolsar (refund, con confirmación JS)"],
        ["rejected", "Mostrar motivo de rechazo"],
        ["Otros", "Sin acciones disponibles"],
    ],
)

add_screenshot("20_custom_admin_payments", "Figura 17: Gestión de pagos y transferencias")

# 4.6 Eventos Panel
doc.add_heading("4.6. Torneos, Eventos y Noticias", level=2)
doc.add_paragraph(
    "Gestión centralizada de todo el contenido del complejo: torneos, "
    "eventos especiales y noticias para los usuarios."
)

doc.add_heading("Crear un Torneo:", level=3)
doc.add_paragraph(
    "1. Haga clic en \"Nuevo Torneo\"\n"
    "2. Ingrese el título del torneo\n"
    "3. Seleccione la categoría (Primera, Segunda, Tercera, Open, Mixto)\n"
    "4. Defina el máximo de parejas/equipos\n"
    "5. Establezca la cuota de inscripción ($)\n"
    "6. Seleccione fecha de inicio y fin\n"
    "7. Haga clic en \"Crear torneo\"\n"
    "8. El torneo se crea con estado 'open'"
)

doc.add_heading("Publicar una Noticia:", level=3)
doc.add_paragraph(
    "1. Haga clic en \"Nueva Noticia\"\n"
    "2. Ingrese el título\n"
    "3. Escriba el contenido\n"
    "4. Haga clic en \"Publicar\"\n"
    "5. Se envían notificaciones a todos los usuarios"
)

add_screenshot("21_custom_admin_events", "Figura 18: Torneos, eventos y noticias")

# 4.7 Reportes
doc.add_heading("4.7. Analítica de Negocio y Reportes", level=2)
doc.add_paragraph(
    "Sección de reportes para el análisis del rendimiento del complejo. "
    "Utiliza el servicio ReportService (apps/reports/services.py) para "
    "generar métricas y exportaciones."
)

doc.add_heading("KPI Principal:", level=3)
add_bullet("Ingresos Mensuales Acumulados — Total de pagos confirmados en el mes actual ($)")

doc.add_heading("Reportes disponibles:", level=3)
add_bullet("Reservas por Estado — Distribución de reservas según su estado (count)")
add_bullet("Ingresos por Cancha — Desglose de ingresos por cada cancha (total $)")
add_bullet("Top 10 Clientes Frecuentes — Los 10 usuarios con más reservas (email, nombre, count)")

doc.add_heading("Exportar CSV:", level=3)
add_bullet("Botón \"Exportar CSV\" descarga un reporte completo")
add_bullet("Formato: andes_padel_report.csv")
add_bullet("Contenido: últimas 500 reservas con ID, Fecha, Cliente, Cancha, Precio, Estado")

doc.add_paragraph(
    "Métodos del ReportService:\n"
    "- revenue_by_period(start, end) — Ingresos por período\n"
    "- revenue_by_court(start, end) — Ingresos por cancha\n"
    "- occupancy_percentage(start, end) — Porcentaje de ocupación\n"
    "- top_customers(start, end, limit) — Top clientes por ingresos\n"
    "- cancellation_rate(start, end) — Tasa de cancelación"
)

add_screenshot("22_custom_admin_reports", "Figura 19: Analítica de negocio y reportes")

# 4.8 Configuración
doc.add_heading("4.8. Configuración (Políticas y Tarifas)", level=2)
doc.add_paragraph(
    "Vista de solo lectura que muestra las políticas de cancelación y "
    "reglas de precios actualmente activas en el sistema."
)

doc.add_heading("Política de Cancelación:", level=3)
add_bullet("Ventana Gratuita (horas) — Horas mínimas antes de la reserva para cancelar sin cargo")
add_bullet("Penalización Cancelación Tardía (%) — Porcentaje cobrado si cancela después de la ventana")
add_bullet("Penalización No-Show (%) — Porcentaje cobrado si el cliente no se presenta")
add_bullet("Tiempo de Hold (minutos) — Duración máxima de un slot en hold")

doc.add_heading("Reglas de Tarifa por Hora:", level=3)
add_bullet("Cancha — Nombre de la cancha")
add_bullet("Rango Horario — Hora inicio - Hora fin")
add_bullet("Día de la Semana — Lunes a Domingo")
add_bullet("Tarifa/Hora — Precio por hora en la franja")

doc.add_paragraph(
    "Si no hay políticas configuradas, se muestran valores por defecto:\n"
    "- Cancelación gratuita: ≥ 24 horas antes\n"
    "- Penalización tardía: 50% del valor\n"
    "- No-show: 100% del valor"
)

add_screenshot("23_custom_admin_settings", "Figura 20: Configuración del sistema")

# 4.9 Auditoría Panel
doc.add_heading("4.9. Registro de Auditoría y Seguridad", level=2)
doc.add_paragraph(
    "Vista detallada del registro de auditoría con opciones de filtrado "
    "y paginación. Los registros son de solo lectura."
)

doc.add_heading("Filtros de auditoría:", level=3)
add_bullet("Por tipo de acción (action) — dropdown dinámico con valores distintos")
add_bullet("Por entidad (entity) — dropdown dinámico con valores distintos")
add_bullet("Por usuario — Búsqueda por correo electrónico (icontains)")

doc.add_heading("Columnas de la tabla:", level=3)
add_bullet("Fecha (UTC) — Formato YYYY-MM-DD HH:MM:SS")
add_bullet("Usuario — Email o \"Sistema / Anónimo\"")
add_bullet("Acción — Badge monospace (ej: admin.login, booking_create)")
add_bullet("Entidad — Badge (ej: user, booking, court, payment)")
add_bullet("ID Entidad — Truncado, monospace")
add_bullet("IP Origen — Dirección IP del cliente")

doc.add_paragraph("Paginación: 50 registros por página con navegación Previous/Next.")

add_screenshot("24_custom_admin_audit", "Figura 21: Registro de auditoría y seguridad")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  5. DOCUMENTACIÓN API
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Documentación de la API (Swagger / ReDoc)", level=1)

doc.add_paragraph(
    "La API REST del sistema está documentada interactivamente utilizando "
    "drf-spectacular (OpenAPI 3.0). Swagger UI permite explorar y probar "
    "todos los endpoints disponibles, mientras que ReDoc ofrece una vista "
    "alternativa más detallada."
)

doc.add_heading("Acceso a la documentación:", level=2)
add_bullet("Swagger UI — http://[IP]:8000/api/docs/")
add_bullet("ReDoc — http://[IP]:8000/api/redoc/")
add_bullet("Schema (JSON/YAML) — http://[IP]:8000/api/schema/")

doc.add_heading("5.1. Endpoints Disponibles", level=2)

doc.add_heading("Grupos de endpoints principales:", level=3)
add_table(
    ["Grupo", "Endpoints", "Descripción"],
    [
        ["Autenticación", "login, register, verify, refresh, logout", "Sesiones JWT"],
        ["Usuarios", "me, password-reset", "Perfil y recuperación"],
        ["Canchas", "list, create, retrieve, update", "CRUD de canchas"],
        ["Horarios", "availability (por cancha y fecha)", "Slots disponibles"],
        ["Reservas", "list, create, preview, cancel, confirm", "Gestión completa"],
        ["Pagos", "create, confirm, refund", "Procesamiento de pagos"],
        ["Eventos", "list, create", "Torneos y eventos"],
        ["Noticias", "list, create", "Publicaciones"],
        ["Notificaciones", "list, mark-read, preferences", "Notificaciones in-app"],
        ["Reportes", "revenue, occupancy, export", "Analítica y CSV"],
    ],
)

add_screenshot("01_swagger_overview", "Figura 22: Documentación Swagger (vista pública)")
add_screenshot("02_redoc_overview", "Figura 23: Documentación ReDoc")

doc.add_heading("5.2. Autenticación JWT", level=2)
doc.add_paragraph(
    "Para acceder a los endpoints protegidos, es necesario autenticarse "
    "utilizando tokens JWT (JSON Web Token). El proceso es:"
)

doc.add_paragraph(
    "1. Realizar una solicitud POST a /api/auth/login/ con email y contraseña\n"
    "2. El servidor responde con un token de acceso (access) y uno de refresco (refresh)\n"
    "3. Incluir el header Authorization: Bearer <access_token> en las solicitudes\n"
    "4. Cuando el access token expire, usar el refresh token para obtener uno nuevo\n"
    "5. Los tokens de refresco tienen rotación y blacklist para seguridad"
)

doc.add_paragraph(
    "En Swagger UI, puede hacer clic en el botón \"Authorize\" (candado) "
    "e ingresar su token de acceso para probar los endpoints protegidos "
    "directamente desde la interfaz."
)

add_screenshot("25_swagger_authenticated", "Figura 24: Swagger con autenticación JWT activa")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  6. CUENTAS DE PRUEBA
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("6. Cuentas de Prueba", level=1)

doc.add_paragraph(
    "Las siguientes cuentas están disponibles para pruebas en el entorno "
    "de desarrollo. Todas comparten la misma contraseña."
)

add_table(
    ["Cuenta", "Correo Electrónico", "Contraseña", "Rol", "Panel"],
    [
        ["Super Administrador", "admin@andespadel.com", "Andes12345!", "superadmin", "Django + Personalizado"],
        ["Cliente", "cliente@andespadel.com", "Andes12345!", "client", "Solo App Móvil"],
    ],
)

doc.add_paragraph()
doc.add_paragraph(
    "Nota: Estas credenciales son solo para entornos de desarrollo. "
    "En producción, cada usuario debe tener credenciales únicas y seguras. "
    "Las contraseñas deben cumplir con las políticas de seguridad del sistema."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
#  7. SOLUCIÓN DE PROBLEMAS
# ════════════════════════════════════════════════════════════════════════
doc.add_heading("7. Solución de Problemas", level=1)

problems = [
    ("No puedo iniciar sesión al panel personalizado",
     "Verifique que está usando el correo electrónico correcto (no un nombre de usuario). "
     "La contraseña es sensible a mayúsculas y minúsculas. Solo usuarios con rol de personal "
     "(recepcionista, gerente, dueño, superadmin) pueden acceder al panel. "
     "Si olvidó su contraseña, contacte al administrador."),
    ("La página no carga o muestra error 500",
     "El servidor puede estar reiniciándose. Espere unos segundos y recargue la página. "
     "Si el problema persiste, verifique que los contenedores Docker estén ejecutándose "
     "con el comando: docker compose ps. Asegúrese de que el backend, db y redis estén activos."),
    ("No se guardan los cambios en el panel Django",
     "Verifique que tiene permisos de edición para la sección. El control de acceso RBAC "
     "limita los permisos según su rol. Los módulos financieros (Payment, CancellationPolicy) "
     "solo son editables por superadmin y dueño."),
    ("El calendario no muestra las canchas",
     "Asegúrese de que existan canchas activas en el sistema. Las canchas inactivas no "
     "aparecen en el calendario. Cree canchas desde Gestión de Canchas o active las existentes."),
    ("Los pagos por transferencia no se confirman",
     "Los pagos por transferencia requieren verificación manual. Vaya a Pagos y Verificación, "
     "encuentre el pago con estado 'pending_transfer' y haga clic en 'Aprobar'. "
     "Puede verificar el comprobante de pago antes de aprobar."),
    ("Las reservas manuales no aparecen en el calendario",
     "Las reservas manuales se crean con estado 'confirmed'. Verifique que la fecha "
     "seleccionada en el calendario coincida con la fecha de la reserva creada."),
    ("No se envían notificaciones",
     "Las notificaciones in-app y por email se envían automáticamente. Las notificaciones "
     "push requieren configuración de Firebase Cloud Messaging. Verifique que Celery esté "
     "ejecutándose: docker compose ps worker."),
    ("El exportar CSV no funciona",
     "Verifique que tenga conexión a la base de datos. El archivo CSV se genera con "
     "las últimas 500 reservas. Si no hay reservas, el archivo estará vacío."),
]

for problem, solution in problems:
    doc.add_heading(problem, level=3)
    doc.add_paragraph(solution)

doc.add_paragraph()
doc.add_paragraph()

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— Fin del Manual —")
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = RGBColor(0x00, 0x2F, 0x48)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Andes Pádel © 2026 — Todos los derechos reservados")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ──
doc.save(str(OUTPUT))
print(f"Manual saved: {OUTPUT}")
print(f"Sections: 7 chapters, 24 screenshots embedded")
